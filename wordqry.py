# -*- coding: utf-8 -*- 
from docx import Document
from docx.shared import Inches
import os
import sys
from unrar import rarfile
import shutil
import subprocess
import openpyxl
from openpyxl import load_workbook
import win32com.client as win32
import docx
import xlrd
import win32api
import win32con

reload(sys)
sys.setdefaultencoding( "utf-8" )

class wordqry:
# 检索word文件并复制到新的word文件
    def FindCopyToWord(self, olePath, newPath, copyData):
        print "asd"
        findDt = self.testWd(olePath, copyData)
        self.saves(newPath, str(findDt))
#查找word文档中的某个字段
    def testWd(self,selfPath,selfData):
        textdoc = ''

        fls = Document(str(selfPath))
        for docxx in fls.paragraphs:
            print str(selfData)
            if str(docxx.text).find(str(selfData)) != -1:
                textdoc = str(docxx.text)
                print(docxx.text)
        return textdoc
#读取word文档
    def readWd(self,selfPath) :
        if os.path.exists(selfPath):
            filesWd = Document(selfPath)
            return filesWd
        else:
            return ''
    def saves(self,selfPath, selpathData):
        if not os.path.exists(selfPath):
            # 创建内存中的word文档对象
            file = Document()
            p = file.add_paragraph(str(selpathData))
        else:
            file = self.readWd(selfPath)
            file.add_paragraph(str(selpathData))
        file.save(selfPath)
#解压缩rar
    def rartest(self,filepath,expath):
        file = rarfile.RarFile(filepath)
        file.extractall(expath)
#移动文件夹
    def filemove(self,oldpath,newpath):
        shutil.move(oldpath,newpath)
#调用bat
    def openbat(self,batpath):
        p = subprocess.Popen("cmd.exe /c" + batpath, stdout=subprocess.PIPE,stderr=subprocess.STDOUT)
        curline = p.stdout.readline()
        while (curline != b''):
            print(curline)
            curline = p.stdout.readline()
        p.wait()
        return p.returncode
#打开excel
    def openexcel(self,excelpath):
        os.startfile(excelpath)
#excel读写
    def exceltest(self,excelpath,sheetname,colname,colvalue,savepath):
        wb = load_workbook(excelpath)
        sheet = wb.get_sheet_by_name(sheetname)
        sheet[colname].value = colvalue
        wb.save(savepath)
#根据配置excel用outlook发送邮件
    def emailsend(self,workbookpath):
        readbook = xlrd.open_workbook(workbookpath)
        sheet = readbook.sheet_by_index(0)
        inputReceivers = str(sheet.cell(1, 0).value)
        inputSubject = str(sheet.cell(1, 1).value)
        docPath = str(sheet.cell(1, 2).value)
        Attachments = str(sheet.cell(1, 3).value)
        doc = docx.Document(docPath)
        doc1 = doc.paragraphs
        b = ''
        for i in doc1:
            b = b + i.text
        s = []
        s.append(inputReceivers)
        print (inputReceivers, inputSubject, docPath, Attachments)
        print (b)
        print (s)
        outlook = win32.Dispatch('outlook.application')
        receivers = s
        mail = outlook.CreateItem(0)
        mail.To = receivers[0]
        mail.Subject = inputSubject
        mail.Body = b
        mail.Attachments.Add(Attachments)
        mail.Send()
#另存为excel
    def excelsave(self,excelpath,savepath):
        wb = load_workbook(excelpath)
        wb.save(savepath)
#工作台提示
    def tip(self,word):
        print (word)
#模拟键盘录入---ctrl+alt+F5
    def ctrlaltf5(self):
        win32api.keybd_event(17, 0, 0, 0)  # ctrl键位码是17
        win32api.keybd_event(18, 0, 0, 0)  # alt键位码是18
        win32api.keybd_event(116, 0, 0, 0)  # F5键位码是116
        win32api.keybd_event(17, 0, win32con.KEYEVENTF_KEYUP, 0)  # 释放按键
        win32api.keybd_event(18, 0, win32con.KEYEVENTF_KEYUP, 0)
        win32api.keybd_event(116, 0, win32con.KEYEVENTF_KEYUP, 0)